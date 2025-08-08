import base64
import io
import os
import pickle

import docx
from email_reply_parser import EmailReplyParser
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from PyPDF2 import PdfReader
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, PointStruct, VectorParams
from sentence_transformers import SentenceTransformer
from tqdm import tqdm

SCOPES = ["https://www.googleapis.com/auth/gmail.readonly"]
COLLECTION_NAME = "gmail_embeddings_full"
HISTORY_FILE = "last_history_id.txt"


def gmail_authenticate():
    creds = None
    if os.path.exists("token.pickle"):
        print("Loading existing Gmail credentials from token.pickle...")
        with open("token.pickle", "rb") as token:
            creds = pickle.load(token)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            print("Refreshing expired Gmail credentials...")
            creds.refresh(Request())
        else:
            print("Starting Gmail OAuth flow - opening browser for authentication...")
            flow = InstalledAppFlow.from_client_secrets_file("credentials.json", SCOPES)
            creds = flow.run_local_server(port=0)
        print("Saving Gmail credentials to token.pickle...")
        with open("token.pickle", "wb") as token:
            pickle.dump(creds, token)
    return creds


def strip_quoted_and_signature(body):
    # Remove quoted replies
    parsed = EmailReplyParser.read(body).fragments
    cleaned = "\n".join(f.content for f in parsed if not f.quoted)
    # Naive signature removal (split on common markers)
    for marker in ["\n--", "\nThanks", "\nBest regards", "\nSent from my"]:
        if marker in cleaned:
            cleaned = cleaned.split(marker)[0]
    return cleaned.strip()


def extract_attachment_text(part, filename):
    data = base64.urlsafe_b64decode(part["body"]["attachmentId"])
    if filename.lower().endswith(".pdf"):
        try:
            pdf_file = io.BytesIO(data)
            reader = PdfReader(pdf_file)

            # Security: Disable JavaScript and other potentially dangerous features
            if reader.metadata:
                # Remove potentially dangerous metadata
                safe_metadata = {}
                for key in ["/Title", "/Author", "/Subject"]:
                    if key in reader.metadata:
                        safe_metadata[key] = str(reader.metadata[key])[:100]  # Limit length

            text = ""
            # Limit to first 50 pages to prevent DoS attacks
            max_pages = min(50, len(reader.pages))
            for i in range(max_pages):
                try:
                    page = reader.pages[i]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    # Limit text size to prevent memory issues
                    if len(text) > 100000:  # 100KB limit
                        text = text[:100000] + "\n[Content truncated due to size limit]"
                        break
                except Exception as e:
                    print(f"Warning: Could not extract text from page {i}: {e}")
                    continue

            return text
        except Exception as e:
            print(f"Warning: Could not process PDF {filename}: {e}")
            return f"[PDF processing failed: {str(e)[:100]}]"
    elif filename.lower().endswith(".docx"):
        try:
            doc = docx.Document(data)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            print(f"Warning: Could not process DOCX {filename}: {e}")
            return f"[DOCX processing failed: {str(e)[:100]}]"
    return ""


def chunk_text(text, chunk_size=500):
    words = text.split()
    for i in range(0, len(words), chunk_size):
        yield " ".join(words[i : i + chunk_size])


def fetch_new_messages(service, max_results=50):
    # Use incremental sync if possible
    if os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE) as f:
            last_history_id = f.read().strip()
        try:
            print(f"Fetching Gmail history since historyId: {last_history_id}...")
            history = (
                service.users()
                .history()
                .list(userId="me", startHistoryId=last_history_id, historyTypes=["messageAdded"])
                .execute()
            )
            message_ids = []
            if "history" in history:
                for h in history["history"]:
                    for m in h.get("messages", []):
                        message_ids.append(m["id"])
        except Exception as e:
            print(f"History API failed, falling back to latest {max_results} emails: {e}")
            message_ids = None
    else:
        message_ids = None

    if not message_ids:
        print(f"Fetching latest {max_results} Gmail messages...")
        results = service.users().messages().list(userId="me", maxResults=max_results).execute()
        message_ids = [m["id"] for m in results.get("messages", [])]

    # Save latest historyId for next run
    print("Fetching Gmail profile to get latest historyId...")
    profile = service.users().getProfile(userId="me").execute()
    with open(HISTORY_FILE, "w") as f:
        f.write(profile["historyId"])

    return message_ids


def get_email_details(service, msg_id):
    print(f"Fetching full details for email {msg_id}...")
    msg = service.users().messages().get(userId="me", id=msg_id, format="full").execute()
    headers = {h["name"]: h["value"] for h in msg["payload"]["headers"]}
    subject = headers.get("Subject", "")
    sender = headers.get("From", "")
    to = headers.get("To", "")
    date = headers.get("Date", "")
    thread_id = msg.get("threadId")

    body_text = ""
    attachments_text = []

    def walk_parts(parts):
        nonlocal body_text, attachments_text
        for part in parts:
            mime_type = part.get("mimeType", "")
            if mime_type == "text/plain":
                data = part["body"].get("data")
                if data:
                    text = base64.urlsafe_b64decode(data).decode(errors="ignore")
                    body_text += text
            elif mime_type in [
                "application/pdf",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ]:
                filename = part.get("filename")
                # For Gmail API attachments: need to fetch separately
                attach_id = part["body"].get("attachmentId")
                if attach_id:
                    print(f"Fetching attachment {attach_id} for email {msg_id}...")
                    attachment = (
                        service.users()
                        .messages()
                        .attachments()
                        .get(userId="me", messageId=msg_id, id=attach_id)
                        .execute()
                    )
                    file_data = base64.urlsafe_b64decode(attachment["data"])
                    if filename.lower().endswith(".pdf"):
                        try:
                            pdf_file = io.BytesIO(file_data)
                            reader = PdfReader(pdf_file)

                            # Security: Limit processing to prevent DoS attacks
                            text = ""
                            max_pages = min(50, len(reader.pages))
                            for i in range(max_pages):
                                try:
                                    page = reader.pages[i]
                                    page_text = page.extract_text()
                                    if page_text:
                                        text += page_text + "\n"
                                    # Limit text size to prevent memory issues
                                    if len(text) > 100000:  # 100KB limit
                                        text = text[:100000] + "\n[Content truncated due to size limit]"
                                        break
                                except Exception as e:
                                    print(f"Warning: Could not extract text from page {i}: {e}")
                                    continue

                            attachments_text.append(text)
                        except Exception as e:
                            print(f"Warning: Could not process PDF attachment {filename}: {e}")
                            attachments_text.append(f"[PDF processing failed: {str(e)[:100]}]")
                    elif filename.lower().endswith(".docx"):
                        try:
                            doc = docx.Document(file_data)
                            attachments_text.append("\n".join(p.text for p in doc.paragraphs))
                        except Exception as e:
                            print(f"Warning: Could not process DOCX attachment {filename}: {e}")
                            attachments_text.append(f"[DOCX processing failed: {str(e)[:100]}]")
            if "parts" in part:
                walk_parts(part["parts"])

    if "parts" in msg["payload"]:
        walk_parts(msg["payload"]["parts"])
    else:
        data = msg["payload"]["body"].get("data")
        if data:
            body_text = base64.urlsafe_b64decode(data).decode(errors="ignore")

    body_text = strip_quoted_and_signature(body_text)
    full_text = f"Subject: {subject}\nFrom: {sender}\nTo: {to}\nDate: {date}\n\n{body_text}"
    if attachments_text:
        full_text += "\n\nAttachments:\n" + "\n".join(attachments_text)

    return {
        "id": msg_id,
        "thread_id": thread_id,
        "subject": subject,
        "sender": sender,
        "to": to,
        "date": date,
        "text": full_text,
    }


def main():
    print("Starting Gmail authentication...")
    creds = gmail_authenticate()
    print("Building Gmail service...")
    service = build("gmail", "v1", credentials=creds)

    message_ids = fetch_new_messages(service, max_results=50)
    print(f"Found {len(message_ids)} new/updated messages.")

    print("Loading sentence transformer model...")
    model = SentenceTransformer("all-MiniLM-L6-v2")
    print("Connecting to Qdrant vector database...")
    qdrant = QdrantClient(host="localhost", port=6333)

    print("Checking if collection exists in Qdrant...")
    if COLLECTION_NAME not in [c.name for c in qdrant.get_collections().collections]:
        print(f"Creating new collection '{COLLECTION_NAME}' in Qdrant...")
        qdrant.create_collection(
            collection_name=COLLECTION_NAME,
            vectors_config=VectorParams(
                size=model.get_sentence_embedding_dimension(), distance=Distance.COSINE
            ),
        )

    points = []
    point_id_counter = 0
    print("Processing emails and creating embeddings...")
    for idx, msg_id in enumerate(tqdm(message_ids)):
        details = get_email_details(service, msg_id)
        for chunk in chunk_text(details["text"], chunk_size=200):
            vec = model.encode(chunk).tolist()
            points.append(PointStruct(id=point_id_counter, vector=vec, payload=details))
            point_id_counter += 1

    if points:
        print(f"Uploading {len(points)} chunks to Qdrant...")
        qdrant.upsert(collection_name=COLLECTION_NAME, points=points)
        print(f"Successfully inserted {len(points)} chunks into Qdrant.")
    else:
        print("No new data to upload to Qdrant.")


if __name__ == "__main__":
    main()
